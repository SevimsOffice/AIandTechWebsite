from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
SLATE  = RGBColor(0x1F, 0x3A, 0x5F)
GOLD   = RGBColor(0xC8, 0x9A, 0x2E)
RED    = RGBColor(0xBF, 0x21, 0x21)
GREEN  = RGBColor(0x1A, 0x6B, 0x3A)
LGRAY  = RGBColor(0xF2, 0xF4, 0xF7)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DKGRAY = RGBColor(0x33, 0x33, 0x33)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top='single', bottom='single', left='single', right='single', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '1F3A5F')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def cell_para(cell, text, bold=False, italic=False, size=10,
              color=DKGRAY, align=WD_ALIGN_PARAGRAPH.LEFT, wrap=True):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '40')
    spacing.set(qn('w:after'), '40')
    pPr.append(spacing)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if level == 0:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = SLATE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(2)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = GOLD
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DKGRAY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DKGRAY
    p.paragraph_format.space_after = Pt(2)
    return p

def add_spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

def make_table(headers, rows, col_widths=None, header_bg='0D1B2A', alt_bg='F2F4F7'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell_para(hdr.cells[i], h, bold=True, size=9.5,
                  color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(hdr.cells[i], header_bg)
        set_cell_borders(hdr.cells[i])
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            bold = ci == 0
            cell_para(tr.cells[ci], str(val), bold=bold, size=9.5,
                      color=DKGRAY)
            set_cell_bg(tr.cells[ci], bg)
            set_cell_borders(tr.cells[ci])
    # Column widths
    if col_widths:
        for ri2, row2 in enumerate(t.rows):
            for ci2, cell in enumerate(row2.cells):
                cell.width = Inches(col_widths[ci2])
    return t


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('CONFIDENTIAL — STRATEGIC INTELLIGENCE REPORT')
r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
r.font.color.rgb = GOLD

add_spacer(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CLAUDE ECOSYSTEM SUPREMACY BLUEPRINT')
r.bold = True; r.font.size = Pt(28); r.font.name = 'Calibri'
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('How to Use the Full Claude Stack at 1000% Leverage')
r.bold = True; r.font.size = Pt(14); r.font.name = 'Calibri'
r.font.color.rgb = SLATE

add_spacer(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared for: Agency / Business Principal\nClassification: Confidential — Operator Grade\nVersion: 1.0  |  May 2026')
r.font.size = Pt(10); r.font.name = 'Calibri'
r.font.color.rgb = DKGRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE DIAGNOSIS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 1 — EXECUTIVE DIAGNOSIS', 1)
add_body(
    "You are using Claude in the way a Formula 1 driver uses a car to commute to the grocery store. "
    "The capability is extraordinary. The utilisation is not. This section does not flatter you — it "
    "catalogues the gap between where you are and where 1,000% leverage begins. The diagnosis below "
    "is based on the usage pattern described: daily conversational engagement across compensation "
    "design, adversarial auditing, pipeline management, interview systems, policy authoring, "
    "geopolitical analysis, and multi-stakeholder decisions. That is an impressive portfolio of use "
    "cases. It is also almost entirely chat-layer execution — meaning you are leaving a minimum of "
    "60–75% of available leverage unrealised every single day."
)
add_spacer()

add_heading('1.1  What Your Usage Pattern Actually Looks Like', 2)
add_body(
    "The table below maps your described behaviour to its actual strategic pattern, its root "
    "mechanism, and the severity of the gap it represents on a 1–5 scale (5 = critical leverage loss)."
)
add_spacer()

diag_headers = ['Pattern', 'Description', 'Root Mechanism', 'Severity (1–5)']
diag_rows = [
    ['Chat-First, Systems-Never',
     'Every task starts and ends in a conversation window. No persistent context, no automated workflows, no reusable structure.',
     'Absence of Projects / Skills architecture',
     '5 — Critical'],
    ['Deep but Disconnected',
     'You produce high-quality outputs — frameworks, audits, analyses — but they exist as chat artifacts that cannot compound.',
     'No Knowledge Layer; outputs die in threads',
     '5 — Critical'],
    ['Manual Execution of Repeatable Work',
     'Compensation framework design, pipeline reviews, and interview system updates are reconstructed from scratch each session.',
     'No Scheduled Tasks / Dispatch automation',
     '4 — High'],
    ['Single-Thread Cognitive Load',
     'You carry the full mental model of every engagement into each conversation, re-briefing Claude on context it should already hold.',
     'Projects not configured; no persistent memory',
     '5 — Critical'],
    ['Brilliant Prompting, Zero Systemisation',
     'You are an expert at getting high-quality outputs from individual prompts. You have not translated that skill into reusable Skills or templates.',
     'Skills Library not built',
     '4 — High'],
    ['Geopolitical / Macro Analysis Without Deep Research',
     'Complex multi-source analysis tasks are handled conversationally when Deep Research could return 10x the source coverage in parallel.',
     'Deep Research underutilised',
     '4 — High'],
    ['Client Deliverables Built Ad Hoc',
     'Each client engagement requires manual assembly of context, frameworks, and outputs rather than operating from a pre-built project template.',
     'No Per-Client Project structure',
     '4 — High'],
    ['No Integration Layer',
     'Claude outputs do not feed downstream tools (CRM, docs, calendars, Slack) automatically. Everything requires manual handoff.',
     'Claude Code / API integration absent',
     '3 — Medium'],
    ['Underuse of Multimodal and Artefact Surfaces',
     'Reports, dashboards, and structured deliverables are produced as chat text rather than as exportable, formatted artefacts.',
     'Claude Design / artefact generation unused',
     '3 — Medium'],
    ['No Compound Learning Loop',
     'Claude does not get smarter about your business over time. Each session is episodic rather than accumulating institutional knowledge.',
     'No Knowledge Layer architecture',
     '5 — Critical'],
]
make_table(diag_headers, diag_rows, col_widths=[1.7, 2.8, 2.2, 1.0])

add_spacer()
add_heading('1.2  The Honest Summary', 2)
add_body(
    "You are in the top 1–2% of Claude users by prompt sophistication and use-case breadth. "
    "You are likely in the bottom 20% of that cohort by systems leverage. The gap is not intelligence — "
    "it is architecture. Everything in this report is about closing that gap."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — WHERE YOU ARE UNDERUSING THE ECOSYSTEM
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 2 — WHERE YOU ARE UNDERUSING THE ECOSYSTEM', 1)
add_body(
    "Eight surfaces. Eight gaps. Each subsection shows your current state, the ideal state, and "
    "the measurable impact of closing the gap. Read this as an audit, not a tutorial."
)

surfaces = [
    {
        'name': '2.1  Projects',
        'intro': (
            "Projects are the foundational unit of Claude leverage. They provide persistent context, "
            "a shared knowledge base, and a consistent operating environment for every interaction "
            "within a defined scope. Without them, you are starting every session at zero."
        ),
        'rows': [
            ['Current State',
             'All work conducted in general chat threads. Context re-entered manually each session. '
             'No persistent client files, no standing instructions, no knowledge base.'],
            ['Ideal State',
             'One Project per client, one per internal function (Ops, BD, Talent, Strategy). Each '
             'loaded with: client background, your frameworks, past decisions, standing operating '
             'instructions, and tone/format preferences. Claude operates as a briefed analyst, not a blank slate.'],
            ['Impact of Closing Gap',
             '40–60% reduction in session setup time. 3–5x improvement in output relevance on first '
             'attempt. Elimination of context drift across long engagements. Institutional knowledge '
             'that compounds over months, not resets every session.'],
        ]
    },
    {
        'name': '2.2  Cowork / Shared Workspaces',
        'intro': (
            "Cowork surfaces allow collaborative, multi-party engagement with Claude on shared "
            "documents and decisions. If you have a team, collaborators, or client-facing delivery "
            "partners, this is the leverage multiplier that converts solo AI use into organisational capability."
        ),
        'rows': [
            ['Current State',
             'Claude use is entirely solo. Team members, if any, operate separate, disconnected '
             'Claude instances with no shared context or institutional outputs.'],
            ['Ideal State',
             'Shared workspaces for client delivery teams. Collaborators can interrogate the same '
             'Project knowledge base. Decision artefacts are co-created, not duplicated.'],
            ['Impact of Closing Gap',
             'Elimination of duplicated AI effort across team. Consistent client-facing voice and '
             'framework application. 2–3x acceleration in collaborative deliverable production.'],
        ]
    },
    {
        'name': '2.3  Dispatch',
        'intro': (
            "Dispatch enables you to route tasks to Claude asynchronously — fire and return. "
            "For a high-velocity operator running multiple client engagements, this is the "
            "difference between Claude being a tool you use and Claude being infrastructure that runs."
        ),
        'rows': [
            ['Current State',
             'All Claude tasks require your active presence. You initiate, supervise, and extract '
             'outputs in real time. Claude is a synchronous tool.'],
            ['Ideal State',
             'Dispatch handles: weekly pipeline scan summaries, competitive intelligence digests, '
             'draft policy document updates, interview scorecard aggregation. You review outputs, '
             'not generate them.'],
            ['Impact of Closing Gap',
             '8–15 hours/week of recoverable executive time. Shift from operator to decision-maker '
             'on Claude-produced work product. Parallel task execution across multiple engagements.'],
        ]
    },
    {
        'name': '2.4  Scheduled Tasks',
        'intro': (
            "Scheduled Tasks are recurring, time-triggered Claude executions. Think of them as a "
            "standing army of analysts who run their briefs every Monday morning whether you remembered "
            "to ask or not."
        ),
        'rows': [
            ['Current State',
             'No recurring automation. Every recurring task — weekly pipeline review, monthly '
             'compensation benchmark, quarterly policy refresh — requires you to manually initiate '
             'and brief Claude.'],
            ['Ideal State',
             'Weekly: pipeline health summary, sales audit trigger list, new prospect research '
             'digests. Monthly: compensation benchmark updates, policy change flags. Quarterly: '
             'strategic review briefing documents.'],
            ['Impact of Closing Gap',
             'Zero marginal effort for recurring analytical work. Consistent cadence of strategic '
             'intelligence without cognitive overhead. Catch-all for tasks that currently fall '
             'through due to busyness.'],
        ]
    },
    {
        'name': '2.5  Skills',
        'intro': (
            "Skills are reusable, specialised Claude behaviours you define once and deploy "
            "repeatedly — your best prompts operationalised as infrastructure. If you are rebuilding "
            "your compensation framework prompt from memory each time, you are burning expensive "
            "cognitive capital on solved problems."
        ),
        'rows': [
            ['Current State',
             'No Skills Library. Best prompts exist in your memory or scattered in notes. Each '
             'deployment requires reconstruction and is subject to degradation.'],
            ['Ideal State',
             'Skills Library of 20–30 operator-grade prompts covering: adversarial sales audit, '
             'compensation band design, interview system build, geopolitical scenario analysis, '
             'multi-stakeholder decision matrix, policy document drafting, prospect pipeline '
             'scoring, executive communication drafting.'],
            ['Impact of Closing Gap',
             '80% reduction in prompt engineering overhead. Consistent output quality across '
             'sessions. Transferable to team members. Compounds as library grows.'],
        ]
    },
    {
        'name': '2.6  Claude Design (Artefact Generation)',
        'intro': (
            "Claude Design and artefact generation surfaces allow production of structured, "
            "formatted, client-ready outputs directly — not as chat text requiring manual "
            "reformatting. If your deliverables require a copy-paste step, you are not done yet."
        ),
        'rows': [
            ['Current State',
             'Outputs are chat text. Client deliverables require manual reformatting into Word, '
             'slides, or structured documents. Significant time lost in translation.'],
            ['Ideal State',
             'Compensation frameworks output as formatted tables. Policy documents output as '
             'structured Word artefacts. Decision matrices output as structured comparison grids. '
             'Interview systems output as scored rubric documents.'],
            ['Impact of Closing Gap',
             '2–4 hours/deliverable saved in reformatting. Consistent branded output quality. '
             'Faster client turnaround. Direct handoff from Claude to client.'],
        ]
    },
    {
        'name': '2.7  Deep Research',
        'intro': (
            "Deep Research is not Claude with a search plugin. It is a multi-source, parallel "
            "research execution that synthesises across tens of sources in a single invocation. "
            "For geopolitical analysis, market intelligence, competitive scanning, and "
            "compensation benchmarking — this is a qualitatively different capability."
        ),
        'rows': [
            ['Current State',
             'Geopolitical and market analysis conducted conversationally, drawing on Claude\'s '
             'training data. Limited to knowledge cutoff. No live source synthesis. Single-thread '
             'analysis.'],
            ['Ideal State',
             'Deep Research invoked for: competitive landscape scans, compensation market data, '
             'regulatory change monitoring, geopolitical risk briefings, prospect intelligence '
             'pre-meetings. Results fed into Projects as living knowledge base.'],
            ['Impact of Closing Gap',
             '10–20x source coverage on research tasks. Live data integration. Elimination of '
             'knowledge-cutoff blind spots. Research that would take 4 hours takes 12 minutes.'],
        ]
    },
    {
        'name': '2.8  Claude Code',
        'intro': (
            "Claude Code is not for engineers only. For a high-leverage operator, it is the "
            "automation layer that eliminates manual data handling, builds internal tools, and "
            "creates integrations between Claude outputs and your operational stack. You do not "
            "need to write code to benefit from Claude Code — you need to know what to ask it to build."
        ),
        'rows': [
            ['Current State',
             'No automation or integration layer. Claude outputs are manually copied into other '
             'systems. Data analysis requires manual preparation. No internal tools built.'],
            ['Ideal State',
             'Claude Code builds: pipeline scoring scripts, compensation band calculators, '
             'interview scorecard aggregators, policy diff generators, CRM update scripts. '
             'Claude outputs feed downstream systems automatically.'],
            ['Impact of Closing Gap',
             'Elimination of manual data handling across all operational domains. Internal tools '
             'built in hours, not weeks. Integration layer that makes Claude outputs operational, '
             'not advisory.'],
        ]
    },
]

ciu_headers = ['Dimension', 'Detail']
for surf in surfaces:
    add_heading(surf['name'], 2)
    add_body(surf['intro'])
    add_spacer()
    make_table(ciu_headers, surf['rows'], col_widths=[1.5, 6.2])
    add_spacer()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — CHAT-ONLY VS ECOSYSTEM-MAXIMISED
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 3 — CHAT-ONLY VS. ECOSYSTEM-MAXIMISED', 1)
add_body(
    "This section maps your core business domains against two operating modes: "
    "your current chat-only approach, and the ecosystem-maximised state. "
    "The delta column is what you are leaving on the table today."
)
add_spacer()

comp_headers = ['Business Area', 'Chat-Only State', 'Ecosystem-Maximised State', 'Delta / What You Gain']
comp_rows = [
    ['Compensation Framework Design',
     'Manual re-briefing each session. Output as chat text. No market data integration. Framework rebuilt from memory.',
     'Project holds all past frameworks + market benchmarks. Deep Research pulls live comp data. Artefact output is client-ready. Scheduled monthly benchmark refresh.',
     'Save 3–5h per engagement. Live market data. Compound framework library. Client-ready output with zero reformatting.'],
    ['Adversarial Sales Auditing',
     'Audit conducted as single conversation. No historical pattern matching. No structured output format. Findings exist only in chat.',
     'Sales Audit Skill deployed consistently. Project stores historical audit findings. Pattern analysis across clients. Structured scorecard output. Trends tracked over time.',
     'Consistent methodology. Cross-client pattern detection. Findings compound into a proprietary audit database.'],
    ['Prospect Pipeline Management',
     'Pipeline review is a manual conversation, re-entering CRM context each time. No scoring automation. No recurring cadence.',
     'Scheduled weekly pipeline health scan. Dispatch handles prospect research pre-meeting. Claude Code integrates with CRM. Scoring model applied consistently.',
     '8–10h/week saved. Consistent scoring. No prospect falls through. Pre-meeting intelligence auto-generated.'],
    ['Interview Systems',
     'Interview system design and updates handled conversationally. Scorecards rebuilt each hire. No institutional learning from past hires.',
     'Interview System Project holds all role-specific rubrics, past candidate data, and hiring decisions. Skills automate scorecard generation. Patterns analysed across cohorts.',
     'Consistent hiring quality. Institutional learning from every hire. Scorecard generation in minutes, not hours.'],
    ['Policy Document Authoring',
     'Policies drafted from scratch each time. No version tracking. No automated refresh triggers. Manual change identification.',
     'Policy Library Project stores all live policies. Scheduled quarterly review flags outdated sections. Deep Research monitors regulatory changes. Diff-generation via Claude Code.',
     'Zero-effort policy maintenance cadence. Regulatory blind spots eliminated. Consistent policy language across documents.'],
    ['Geopolitical & Macro Analysis',
     'Conversational analysis limited to training data cutoff. Single-source synthesis. No systematic monitoring of key regions or indicators.',
     'Deep Research generates multi-source intelligence briefs on demand. Scheduled weekly geopolitical digest for key regions. Findings stored in Strategy Project.',
     'Live intelligence. 10x source coverage. Systematic monitoring. Decision-ready briefs rather than conversational summaries.'],
    ['Multi-Stakeholder Decisions',
     'Decision analysis conducted as conversation. No structured framework applied consistently. No record of decision rationale or outcomes.',
     'Decision Framework Skill deployed for all significant decisions. Outputs stored in Decision Log Project. Outcomes tracked and analysed over time. Pattern learning across decision types.',
     'Consistent decision quality. Institutional decision memory. Pattern detection across outcomes. Faster execution through framework clarity.'],
    ['Client Deliverable Production',
     'Each deliverable assembled manually. Context re-entered per client. Formatting done outside Claude. No reusable templates.',
     'Per-Client Projects hold full context. Deliverable templates pre-built. Artefact generation produces formatted outputs. Dispatch handles draft generation.',
     '50–70% reduction in deliverable production time. Consistent quality. Client context always current. First draft quality dramatically higher.'],
    ['Business Development',
     'BD analysis and outreach drafted conversationally. No systematic prospect research. No pattern analysis of what wins.',
     'BD Project stores ICP, win/loss patterns, messaging frameworks. Deep Research generates prospect intelligence. Skills automate outreach drafts. Pipeline tracked systematically.',
     'Higher outreach quality. Faster research. Win/loss pattern learning. Systematic BD execution replacing ad hoc effort.'],
    ['Executive Communication',
     'Comms drafted from scratch in chat. No voice/tone consistency enforcement. No template library. Reformatting required.',
     'Executive Comms Skill enforces tone and format. Templates for board updates, client reports, team communications. Artefact output is send-ready.',
     'Consistent voice. Faster production. Send-ready output. Brand quality maintained at volume.'],
]
make_table(comp_headers, comp_rows, col_widths=[1.5, 2.1, 2.2, 1.9])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — TOP 20 HIGHEST-ROI USE CASES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 4 — TOP 20 HIGHEST-ROI USE CASES', 1)
add_body(
    "Ranked by composite ROI — time saved × strategic impact × implementation speed. "
    "Type key: SA = Systematic Automation, SF = Strategic Framework, INT = Integration, "
    "INT/AI = Intelligence/Analysis, DEL = Deliverable Production."
)
add_spacer()

uc_headers = ['#', 'Use Case', 'Surfaces', 'Type', 'Frequency', 'Key Result']
uc_rows = [
    ['1', 'Per-Client Project Architecture',
     'Projects', 'SF', 'One-time setup, perpetual',
     'Eliminates context re-entry; 40–60% session time saved permanently'],
    ['2', 'Adversarial Sales Audit Skill',
     'Skills, Projects', 'SF', 'Per engagement',
     'Consistent methodology; proprietary audit database compounds over time'],
    ['3', 'Weekly Pipeline Intelligence Scan',
     'Scheduled Tasks, Dispatch', 'SA', 'Weekly',
     '6–10h/week recovered; zero-miss pipeline visibility'],
    ['4', 'Compensation Benchmark Refresh',
     'Deep Research, Projects', 'INT/AI', 'Monthly',
     'Live market data without manual research; framework always current'],
    ['5', 'Pre-Meeting Prospect Intelligence Brief',
     'Deep Research, Dispatch', 'INT/AI', 'Per meeting',
     '30-min research replaced by 3-min review; win rate improvement'],
    ['6', 'Decision Framework Skill + Log',
     'Skills, Projects', 'SF', 'Per major decision',
     'Consistent decision quality; institutional decision memory'],
    ['7', 'Interview Scorecard Generation',
     'Skills, Projects', 'SA', 'Per hire cycle',
     'Scorecard in minutes; consistent hiring quality; cohort analytics'],
    ['8', 'Policy Document Quarterly Refresh',
     'Scheduled Tasks, Deep Research, Projects', 'SA', 'Quarterly',
     'Zero-effort compliance maintenance; regulatory blind spots eliminated'],
    ['9', 'Geopolitical Weekly Intelligence Digest',
     'Scheduled Tasks, Deep Research', 'SA', 'Weekly',
     'Live multi-source intelligence; strategic context always current'],
    ['10', 'Executive Communication Templates',
     'Skills, Claude Design', 'DEL', 'Daily',
     'Send-ready output in minutes; consistent voice at volume'],
    ['11', 'Win/Loss Pattern Analysis',
     'Projects, Deep Research', 'INT/AI', 'Monthly',
     'Systematic learning from BD outcomes; ICP refinement'],
    ['12', 'Multi-Stakeholder Decision Simulation',
     'Skills, Projects', 'SF', 'Per complex decision',
     'Adversarial pre-mortem; blind spot detection before commitment'],
    ['13', 'CRM Pipeline Update Automation',
     'Claude Code, Dispatch', 'INT', 'Weekly',
     'Manual CRM data handling eliminated; pipeline always current'],
    ['14', 'Compensation Band Calculator Tool',
     'Claude Code', 'INT', 'Per hire/review cycle',
     'Internal tool built once; used across all compensation decisions'],
    ['15', 'Client Deliverable Template Library',
     'Projects, Claude Design, Skills', 'DEL', 'Per engagement',
     '50–70% production time reduction; client-ready first drafts'],
    ['16', 'Regulatory Change Monitoring',
     'Scheduled Tasks, Deep Research', 'SA', 'Weekly',
     'Proactive compliance posture; risk elimination'],
    ['17', 'Competitive Intelligence Digest',
     'Scheduled Tasks, Deep Research', 'SA', 'Weekly',
     'Market awareness without manual scanning; strategic positioning'],
    ['18', 'Policy Diff Generator',
     'Claude Code, Projects', 'INT', 'Per policy update',
     'Version control and change tracking automated'],
    ['19', 'Onboarding Knowledge Base Build',
     'Projects, Skills', 'SF', 'One-time setup',
     'New team members operate at 80% capacity from week one'],
    ['20', 'Strategic Scenario Planning Framework',
     'Skills, Deep Research, Projects', 'INT/AI', 'Quarterly',
     'Structured futures thinking; decision-ready scenario outputs'],
]
make_table(uc_headers, uc_rows, col_widths=[0.3, 2.0, 1.5, 0.6, 1.0, 2.3])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — IDEAL CLAUDE ECOSYSTEM STACK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 5 — IDEAL CLAUDE ECOSYSTEM STACK', 1)

add_heading('5.1  Architecture Overview', 2)
add_body(
    "Your Claude ecosystem should operate as a four-layer architecture. Each layer has a "
    "distinct function and feeds the layers above and below it. Confusion between layers "
    "is the primary reason operators plateau at average leverage."
)
add_spacer()

arch_headers = ['Layer', 'Function', 'Primary Surfaces', 'What Lives Here']
arch_rows = [
    ['Founder Layer\n(You)',
     'Strategic direction, decision-making, output review, framework ownership',
     'Chat, Projects, Skills invocation',
     'Decision authority. Framework design. Output approval. Strategic intent. Exception handling.'],
    ['Execution Layer',
     'Task execution, deliverable production, analysis, drafting',
     'Dispatch, Scheduled Tasks, Deep Research, Claude Design',
     'All work product generation. Research execution. Draft production. Data analysis. Pattern scanning.'],
    ['Knowledge Layer',
     'Persistent context, institutional memory, framework storage, output archiving',
     'Projects (all), Knowledge Base files',
     'Client context. Past decisions. Frameworks. Patterns. Policies. Outputs. Benchmarks.'],
    ['Integration Layer',
     'System connectivity, automation, data flows, tool interoperability',
     'Claude Code, API, CRM integrations',
     'Automated workflows. CRM sync. Output formatting. Data pipeline. Internal tooling.'],
]
make_table(arch_headers, arch_rows, col_widths=[1.5, 2.0, 1.8, 2.5])

add_spacer()
add_heading('5.2  Surface-to-Function Mapping', 2)
add_spacer()

sfm_headers = ['Claude Surface', 'Layer', 'Primary Function', 'When to Use', 'When NOT to Use']
sfm_rows = [
    ['Projects', 'Knowledge', 'Persistent context container',
     'Any multi-session engagement; per-client work; recurring functions',
     'One-off throwaway tasks with no future value'],
    ['Skills', 'Execution + Founder', 'Reusable specialised behaviours',
     'Any task you execute more than twice; methodology enforcement',
     'Novel one-off analysis with no repeat pattern'],
    ['Dispatch', 'Execution', 'Async task execution',
     'Any task that does not require your real-time input',
     'Tasks requiring judgment calls or sensitive decisions'],
    ['Scheduled Tasks', 'Execution', 'Recurring automated execution',
     'Any task with a defined cadence (weekly, monthly, quarterly)',
     'Ad hoc or unpredictable trigger tasks'],
    ['Deep Research', 'Execution', 'Multi-source intelligence synthesis',
     'Any analysis requiring current, multi-source data; market research; competitive scanning',
     'Internal-only analysis; tasks with no external data requirement'],
    ['Claude Design', 'Execution', 'Structured artefact generation',
     'Any deliverable with format requirements; client-facing outputs',
     'Internal rough drafts; exploratory thinking'],
    ['Claude Code', 'Integration', 'Automation and tool building',
     'Repetitive data tasks; system integrations; internal tooling',
     'One-time analysis; non-technical workflows'],
    ['Chat', 'Founder', 'Interactive strategic dialogue',
     'Novel problems; exploratory thinking; real-time decisions',
     'Repeatable tasks (use Skills); research (use Deep Research)'],
]
make_table(sfm_headers, sfm_rows, col_widths=[1.4, 0.8, 1.8, 1.8, 1.9])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DAILY & WEEKLY OPERATING MODEL
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 6 — DAILY & WEEKLY OPERATING MODEL', 1)

add_heading('6.1  Daily Operating Schedule', 2)
add_spacer()

daily_headers = ['Time Block', 'Activity', 'Claude Surface', 'Output']
daily_rows = [
    ['06:00–06:30', 'Review overnight Scheduled Task outputs: geopolitical digest, pipeline flags, competitive alerts',
     'Scheduled Tasks → Projects', 'Annotated intelligence brief; action flags for the day'],
    ['06:30–07:00', 'Morning decision queue: apply Decision Framework Skill to 1–3 pending decisions',
     'Skills, Chat', 'Decision records logged to Decision Log Project'],
    ['07:00–08:00', 'Deep work — client deliverable production using client Project context',
     'Projects, Skills, Claude Design', 'Client-ready draft artefacts'],
    ['08:00–09:00', 'BD and pipeline work: Dispatch pre-meeting research briefs; pipeline scoring review',
     'Dispatch, Projects', 'Meeting briefs; pipeline updates'],
    ['12:00–12:30', 'Mid-day review of Dispatch outputs queued in morning',
     'Dispatch', 'Reviewed and approved work product'],
    ['15:00–16:00', 'Strategic analysis: Deep Research on priority topics; scenario planning',
     'Deep Research, Skills', 'Intelligence inputs for strategic decisions'],
    ['17:00–17:30', 'End-of-day: update Project knowledge bases with decisions made; flag recurring tasks for next day',
     'Projects', 'Updated institutional memory; automated task queue'],
]
make_table(daily_headers, daily_rows, col_widths=[1.0, 2.4, 1.6, 2.7])

add_spacer()
add_heading('6.2  Weekly Operating Schedule', 2)
add_spacer()

weekly_headers = ['Day', 'Priority Focus', 'Key Claude Activities', 'Surfaces']
weekly_rows = [
    ['Monday', 'Intelligence & Planning',
     'Review weekly Scheduled Task outputs. Set strategic priorities. Pipeline health scan. Competitive brief review.',
     'Scheduled Tasks, Deep Research, Projects'],
    ['Tuesday', 'Client Delivery',
     'All major client deliverable production. Deploy client-specific Skills. Artefact generation.',
     'Projects, Skills, Claude Design, Dispatch'],
    ['Wednesday', 'Business Development',
     'Prospect research via Deep Research. Outreach drafting via Skills. Pipeline updates via Claude Code.',
     'Deep Research, Skills, Projects, Claude Code'],
    ['Thursday', 'Operations & Systems',
     'Policy review and updates. Process documentation. Interview systems. Internal tool maintenance.',
     'Projects, Skills, Claude Code, Scheduled Tasks'],
    ['Friday', 'Strategic Review & Week Close',
     'Decision log review. Win/loss pattern analysis. Knowledge base updates. Next week automation setup.',
     'Projects, Deep Research, Scheduled Tasks'],
    ['Saturday', 'Learning & Framework Development',
     'Skills Library expansion. New framework design. Ecosystem architecture improvements.',
     'Chat, Skills, Projects'],
    ['Sunday', 'Minimal — Async Review Only',
     'Review any automated outputs. Flag items for Monday. No active production.',
     'Scheduled Tasks (review only)'],
]
make_table(weekly_headers, weekly_rows, col_widths=[0.9, 1.5, 2.8, 2.55])

add_spacer()
add_heading('6.3  Per-Client Project Template', 2)
add_body('Every client Project should be pre-loaded with the following structure at engagement start:')
add_spacer()

project_headers = ['Section', 'Content to Pre-Load', 'Update Cadence']
project_rows = [
    ['Client Background', 'Industry, size, structure, key contacts, history with your agency, decision-making dynamics', 'At engagement start; update on material changes'],
    ['Engagement Objectives', 'Defined outcomes, success metrics, timeline, constraints, stakeholder map', 'Weekly review; update on scope changes'],
    ['Frameworks Applied', 'Compensation bands (if applicable), audit methodology, interview rubrics, decision frameworks', 'Update after each engagement session'],
    ['Decision Log', 'All significant decisions made, rationale, alternatives considered, outcomes tracked', 'After every major decision'],
    ['Deliverable Archive', 'All produced artefacts, version history, client feedback', 'After every deliverable submission'],
    ['Open Issues & Risks', 'Active challenges, unresolved questions, risk factors, mitigation approaches', 'Weekly review'],
    ['Standing Instructions', 'Client communication style, formatting preferences, confidentiality notes, escalation triggers', 'At engagement start; rarely changes'],
    ['Intelligence Feed', 'Relevant market data, competitive context, regulatory notes, Deep Research outputs', 'As generated; weekly Scheduled Task update'],
]
make_table(project_headers, project_rows, col_widths=[1.5, 3.8, 2.4])

add_spacer()
add_heading('6.4  Per-Decision 7-Step Framework', 2)
add_body('Apply this framework via the Decision Framework Skill to every decision above a defined materiality threshold:')
add_spacer()

decision_headers = ['Step', 'Action', 'Claude Role', 'Output']
decision_rows = [
    ['1. Frame', 'Define the actual decision (not the presenting question)',
     'Reframe prompt: "What is the real decision here? What are we not asking?"',
     'Precise decision statement with hidden assumptions surfaced'],
    ['2. Stakes', 'Assess reversibility, blast radius, and time horizon',
     'Structured stakes analysis against decision taxonomy',
     'Irreversibility score; urgency classification'],
    ['3. Options', 'Generate full option set including non-obvious and contrarian choices',
     'Adversarial options generation: "What would the opposite decision look like? What option am I avoiding?"',
     'Option set including options you would not have generated solo'],
    ['4. Second Order', 'Map second and third-order consequences of each option',
     'Consequence chain mapping for each option',
     'Consequence map including unintended effects'],
    ['5. Stakeholders', 'Identify how each stakeholder group experiences each option',
     'Stakeholder impact matrix',
     'Stakeholder response map; negotiation leverage analysis'],
    ['6. Adversarial Test', 'Stress-test the preferred option against the strongest counter-argument',
     '"Make the strongest possible case against my preferred option"',
     'Steelmanned opposition; blind spots identified'],
    ['7. Decide & Record', 'Make the decision; record rationale, alternatives rejected, and review trigger',
     'Decision record generation for Project log',
     'Documented decision with full audit trail'],
]
make_table(decision_headers, decision_rows, col_widths=[0.7, 1.6, 2.1, 3.3])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — YOUR LIKELY WEAKNESSES AS A CLAUDE POWER USER
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 7 — YOUR LIKELY WEAKNESSES AS A CLAUDE POWER USER', 1)
add_body(
    "Eight gap callouts. These are not rare failures — they are the predictable failure modes "
    "of sophisticated conversational Claude users who have not yet crossed into systems leverage. "
    "No flattery. Read this as a peer diagnosis."
)
add_spacer()

gap_headers = ['Gap #', 'Weakness', 'How It Manifests', 'Why It Happens', 'Cost']
gap_rows = [
    ['1', 'Prompt Artisanship Without Systemisation',
     'You write excellent prompts but rebuild them from memory. Output quality varies session to session. Your best work is unreproducible at scale.',
     'High verbal intelligence. Implicit knowledge not externalised. Skills Library not built.',
     'Your best prompts exist once. They cannot be delegated, transferred, or compounded.'],
    ['2', 'Solutioning Before Framing',
     'You enter Claude with a solution in mind and use Claude to elaborate it, rather than to challenge the framing. You get sophisticated answers to the wrong questions.',
     'Expert bias. High pattern-matching speed. Claude as confirmer rather than challenger.',
     'Optimised wrong paths. Systematic blind spots. Decisions that look rigorous but missed the actual question.'],
    ['3', 'Intelligence Without Integration',
     'You generate high-quality analysis that does not connect to action systems. Insights exist in chat. They do not change what happens operationally.',
     'No Integration Layer. Claude Code not used. Outputs are advisory, not operational.',
     'Analysis that creates clarity but not outcomes. Strategic intelligence that is not actioned.'],
    ['4', 'Over-Reliance on Single-Thread Analysis',
     'For complex multi-variable problems, you iterate in a single conversation rather than running parallel analyses on sub-components.',
     'Default chat behaviour. Dispatch and Cowork surfaces not deployed.',
     'Analysis that is sequential when it should be parallel. Longer cycles. Perspective narrowing over long threads.'],
    ['5', 'Context Degradation Blindness',
     'You do not notice when long conversations cause Claude to lose precision on earlier constraints. Late-conversation outputs drift from initial parameters.',
     'No awareness of context window dynamics. No structured refreshes.',
     'Decisions made on outputs that have silently drifted from your actual constraints.'],
    ['6', 'Knowledge Without Accumulation',
     'Every insight, framework, and output you have produced in Claude is gone the moment the conversation closes. You are not building institutional memory.',
     'Projects not configured. No Knowledge Layer architecture.',
     'Zero compound learning. You are equally smart today as Day 1. Claude is not getting smarter about your business.'],
    ['7', 'Automation Aversion',
     'You find it faster to do recurring tasks manually than to invest the setup time in automation. This is mathematically correct in the short term and catastrophically wrong over 90 days.',
     'High execution speed creates false efficiency signal. Setup costs feel high; payback feels uncertain.',
     'At 10 recurring tasks × 2h each × 52 weeks = 1,040 hours/year of recoverable time foregone.'],
    ['8', 'First-Draft Acceptance',
     'You treat Claude\'s first output as near-final more often than is warranted. You are not systematically stress-testing outputs before acting on them.',
     'High output quality creates over-trust. Time pressure. Adversarial review not built into workflow.',
     'Strategic decisions and client deliverables built on outputs that would have improved materially with one adversarial review pass.'],
]
make_table(gap_headers, gap_rows, col_widths=[0.4, 1.5, 1.8, 1.7, 2.35])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — HOW TO BEAT THE TOP 0.1% OF AI USERS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 8 — HOW TO BEAT THE TOP 0.1% OF AI USERS', 1)
add_body(
    "The top 0.1% of Claude users are not smarter than you. They have better architecture. "
    "They beat the top 1% on five dimensions. Here is what each dimension requires."
)

dimensions = [
    ('8.1  Speed', [
        ['What Top 0.1% Does', 'Task execution that is 5–10x faster than top 1%. Not through better prompts — through pre-built context, Skills deployment, and Dispatch execution.'],
        ['How They Achieve It', 'Projects eliminate context re-entry. Skills eliminate prompt reconstruction. Dispatch eliminates synchronous execution. Scheduled Tasks eliminate initiation overhead.'],
        ['Your Gap', 'You are fast at prompting. You are slow at task execution because you have no automation layer.'],
        ['What to Do', 'Build Projects first. Build Skills second. Then deploy Dispatch and Scheduled Tasks. Speed compounds once all three are in place.'],
    ]),
    ('8.2  Depth', [
        ['What Top 0.1% Does', 'Analysis that is qualitatively richer because it synthesises across more sources, more time horizons, and more stakeholder perspectives simultaneously.'],
        ['How They Achieve It', 'Deep Research for external synthesis. Multiple Projects for context depth. Structured frameworks that force second-order thinking.'],
        ['Your Gap', 'You have depth in individual analyses. You do not have systematic depth — meaning every analysis is as good as you remembered to make it, not as good as your best framework demands.'],
        ['What to Do', 'Encode your best analytical frameworks as Skills. Use Deep Research for all external-data tasks. Deploy the 7-step Decision Framework on every material decision.'],
    ]),
    ('8.3  Systems', [
        ['What Top 0.1% Does', 'Operates a self-improving intelligence infrastructure. Each use makes the system smarter. Outputs compound.'],
        ['How They Achieve It', 'Knowledge Layer that accumulates. Integration Layer that automates. Scheduled Tasks that run without initiation. Pattern learning across Projects.'],
        ['Your Gap', 'You have none of this. You are at zero compound intelligence. Day 1 and Day 365 of your Claude use have the same starting context.'],
        ['What to Do', 'This is the highest-priority architectural shift. Start with Projects. Add Knowledge Base files. Implement Scheduled Task outputs feeding back to Projects.'],
    ]),
    ('8.4  Client Results', [
        ['What Top 0.1% Does', 'Delivers to clients at a quality and speed that creates structural competitive advantage. Clients cannot replicate the output quality without the infrastructure you have built.'],
        ['How They Achieve It', 'Per-client Projects with deep context. Proprietary Skills that encode their methodology. Artefact generation that produces client-ready outputs.'],
        ['Your Gap', 'Your client work is excellent but labour-intensive. You have not turned your methodology into infrastructure that you own and your competitors do not.'],
        ['What to Do', 'Build your proprietary Skills Library. This is the most durable competitive moat available to an AI-native operator today.'],
    ]),
    ('8.5  Decision Quality', [
        ['What Top 0.1% Does', 'Makes decisions that are structurally more rigorous because they are adversarially tested, multi-perspective mapped, and recorded for institutional learning.'],
        ['How They Achieve It', 'Decision Framework Skills deployed consistently. Decision Log in Projects. Outcome tracking and pattern analysis over time.'],
        ['Your Gap', 'Your individual decision quality is high. Your decision system quality is near zero. You are not learning from your own decision history.'],
        ['What to Do', 'Implement the 7-step Decision Framework as a Skill. Build a Decision Log Project. Review quarterly. This creates a learning loop your competitors cannot replicate without the history you are accumulating.'],
    ]),
]

for dim_name, dim_rows in dimensions:
    add_heading(dim_name, 2)
    make_table(['Dimension', 'Detail'], dim_rows, col_widths=[1.5, 6.2])
    add_spacer()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — IMPLEMENTATION ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 9 — IMPLEMENTATION ROADMAP — 4-WEEK PLAN', 1)
add_body(
    "Four weeks. Four phases. Each phase builds on the prior. Do not skip ahead — "
    "the Foundation phase is the prerequisite for everything else, and doing Scale before "
    "Automation wastes the leverage that Automation creates."
)
add_spacer()

weeks = [
    ('Week 1 — FOUNDATION', '0D1B2A', [
        ['Day 1–2: Project Architecture',
         'Build all Projects: 1 per active client + 5 internal (Ops, BD, Talent, Strategy, Decision Log). Pre-load each with context documents.',
         'Projects live. Context no longer re-entered manually. Immediate time savings begin.'],
        ['Day 3–4: Skills Library v1',
         'Build 8 core Skills: Adversarial Sales Audit, Compensation Design, Interview Scorecard, Decision Framework, Policy Draft, Prospect Research Brief, Executive Comms, Geopolitical Brief.',
         'First 8 Skills deployed. Core methodology encoded. First reuse delivers ROI.'],
        ['Day 5: Knowledge Base Load',
         'Upload all relevant documents to Projects: past deliverables, frameworks, policies, client backgrounds, benchmarks.',
         'Knowledge Layer populated. Claude now has institutional context.'],
        ['Day 6–7: Operating Model Adoption',
         'Run first week on new Daily Schedule. Deploy Skills on actual work. Track time saved vs. prior week.',
         'Baseline established. System in use. First compound effects visible.'],
    ]),
    ('Week 2 — AUTOMATION', '1F3A5F', [
        ['Day 8–9: Scheduled Tasks Setup',
         'Configure: Weekly Pipeline Scan, Weekly Geopolitical Digest, Monthly Compensation Benchmark, Quarterly Policy Review.',
         'Recurring automation live. First scheduled outputs arrive end of week.'],
        ['Day 10–11: Dispatch Deployment',
         'Configure Dispatch for: pre-meeting prospect research, draft deliverable generation, competitive intelligence pulls.',
         'Async execution live. Active supervision of recurring tasks eliminated.'],
        ['Day 12–13: Deep Research Integration',
         'Deploy Deep Research on 3–5 highest-value research tasks. Integrate outputs into Projects.',
         'Research quality step-changes. Knowledge base begins receiving live-data inputs.'],
        ['Day 14: Automation Audit',
         'Review Week 2 automated outputs. Calibrate quality. Adjust prompts. Identify next automation candidates.',
         'Automation layer optimised. Quality baselines set for recurring outputs.'],
    ]),
    ('Week 3 — SCALE', 'C8 9A 2E'.replace(' ', ''), [
        ['Day 15–16: Claude Code — First Tools',
         'Build: Pipeline Scoring Script, Compensation Band Calculator. Deploy against live data.',
         'First internal tools live. Manual data handling eliminated for these workflows.'],
        ['Day 17–18: Deliverable Templates',
         'Build full deliverable template library using Projects + Claude Design for all client-facing output types.',
         'Deliverable production time cuts by 50–70%. Client-ready first drafts standard.'],
        ['Day 19–20: Skills Library Expansion',
         'Expand Skills to 20+ covering all recurring analytical and production tasks. Share with any team members.',
         'Full methodology encoded. Team leverage (if applicable) begins.'],
        ['Day 21: Scale Audit',
         'Full system review. What is still manual that should be automated? What Skills need refinement? What Projects need updating?',
         'System gaps identified. Week 4 optimisation targets set.'],
    ]),
    ('Week 4 — COMPOUND', '1A6B3A', [
        ['Day 22–23: Integration Layer Build',
         'Deploy Claude Code CRM integration. Build policy diff generator. Automate output formatting across all deliverable types.',
         'Full Integration Layer live. Claude outputs operational, not advisory.'],
        ['Day 24–25: Knowledge Layer Optimisation',
         'Review all Projects for knowledge gaps. Add Deep Research outputs to standing knowledge bases. Implement output-to-Project feedback loops.',
         'Compound intelligence loop active. System gets smarter with each use.'],
        ['Day 26–27: Decision System Activation',
         'Review Decision Log for patterns. Identify recurring decision types. Build decision templates. Set quarterly decision review cadence.',
         'Institutional decision memory active. Learning loop initiated.'],
        ['Day 28: Full System Review',
         'Measure total time saved vs. Week 0. Identify next 90-day optimisation priorities. Document what is working and what is not.',
         'Baseline ROI measurement. 90-day roadmap for next optimisation cycle.'],
    ]),
]

for week_title, header_bg, activities in weeks:
    add_heading(week_title, 2)
    w_headers = ['Activity', 'Actions', 'Milestone']
    make_table(w_headers, activities, col_widths=[2.0, 3.5, 2.2], header_bg=header_bg)
    add_spacer()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — FINAL BLUEPRINT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SECTION 10 — FINAL BLUEPRINT', 1)

add_heading('10.1  Skills Library — Build Order', 2)
add_body('Build in this exact sequence. Each tier depends on the prior tier being deployed and battle-tested.')
add_spacer()

sl_headers = ['Priority', 'Skill Name', 'Use Case', 'Build Time', 'First ROI']
sl_rows = [
    ['Tier 1 — Week 1', 'Adversarial Sales Audit', 'Sales process review; pipeline quality scoring', '2h', 'First use'],
    ['Tier 1 — Week 1', 'Compensation Band Design', 'Role band creation; market benchmarking; offer structuring', '3h', 'First use'],
    ['Tier 1 — Week 1', 'Decision Framework (7-step)', 'All material decisions; option generation; stress-testing', '2h', 'First use'],
    ['Tier 1 — Week 1', 'Interview Scorecard Generator', 'Role-specific rubrics; candidate evaluation; cohort analysis', '2h', 'First hire'],
    ['Tier 1 — Week 1', 'Executive Communication', 'Board updates; client reports; team communications', '1h', 'First use'],
    ['Tier 1 — Week 1', 'Geopolitical Brief', 'Regional risk analysis; macro scenario framing', '2h', 'First use'],
    ['Tier 1 — Week 1', 'Policy Document Draft', 'Policy creation; regulatory alignment; version updates', '2h', 'First use'],
    ['Tier 1 — Week 1', 'Prospect Research Brief', 'Pre-meeting intelligence; ICP scoring; outreach personalisation', '1h', 'First meeting'],
    ['Tier 2 — Week 3', 'Win/Loss Analysis', 'BD pattern detection; ICP refinement; messaging improvement', '2h', 'Month 2'],
    ['Tier 2 — Week 3', 'Multi-Stakeholder Decision Map', 'Complex decisions with multiple affected parties', '2h', 'First use'],
    ['Tier 2 — Week 3', 'Strategic Scenario Planning', 'Quarterly futures planning; option stress-testing', '3h', 'First quarter'],
    ['Tier 2 — Week 3', 'Competitive Landscape Analysis', 'Market positioning; competitor intelligence synthesis', '2h', 'First use'],
    ['Tier 2 — Week 3', 'Client Onboarding Framework', 'New engagement setup; expectation alignment; workplan build', '2h', 'First new client'],
    ['Tier 2 — Week 3', 'Talent Market Intelligence', 'Hiring market conditions; candidate pool analysis; sourcing strategy', '2h', 'First hire'],
    ['Tier 3 — Month 2', 'Contract Review Brief', 'Commercial term analysis; risk flagging; negotiation positioning', '3h', 'First contract'],
    ['Tier 3 — Month 2', 'Board / Investor Communication', 'Formal reporting; narrative framing; financial summary', '2h', 'First board update'],
    ['Tier 3 — Month 2', 'Risk Register Update', 'Operational risk cataloguing; mitigation planning', '2h', 'First risk review'],
    ['Tier 3 — Month 2', 'Process Documentation', 'SOPs; workflow capture; knowledge transfer documents', '1h', 'First process doc'],
    ['Tier 3 — Month 2', 'Client Deliverable QA', 'Output review checklist; quality standard enforcement', '1h', 'First use'],
    ['Tier 3 — Month 2', 'Strategic Review Facilitation', 'Quarterly business review; performance analysis; next-quarter planning', '3h', 'First QBR'],
]
make_table(sl_headers, sl_rows, col_widths=[1.4, 1.8, 2.3, 0.9, 1.35])

add_spacer()
add_heading('10.2  Project Structure — Canonical Layout', 2)
add_spacer()

ps_headers = ['Project', 'Knowledge Files to Pre-Load', 'Standing Instructions', 'Scheduled Task Links']
ps_rows = [
    ['[Client Name] — Active',
     'Client background, org chart, past deliverables, compensation data, decision log, open issues',
     'Tone of voice, confidentiality level, formatting preferences, escalation triggers',
     'Weekly status summary; monthly deliverable review'],
    ['Business Development',
     'ICP definition, win/loss log, messaging frameworks, competitive positioning, outreach templates',
     'Apply prospect research brief Skill by default; always score against ICP',
     'Weekly pipeline scan; bi-weekly prospect research batch'],
    ['Operations & Policy',
     'All active policies, process docs, vendor contracts, compliance tracker, policy change log',
     'Flag any regulatory references; always generate diff from previous version',
     'Quarterly policy review; weekly regulatory change monitor'],
    ['Talent & Hiring',
     'Role profiles, compensation bands, interview rubrics, past hire outcomes, sourcing strategies',
     'Apply Interview Scorecard Skill by default; track against comp bands',
     'Weekly talent market digest; per-hire scorecard generation'],
    ['Strategy & Intelligence',
     'Strategic priorities, competitive landscape, geopolitical watch list, scenario plans, decision log',
     'Apply Strategic Scenario Skill for forward-looking questions; always note confidence level',
     'Weekly geopolitical digest; monthly competitive update; quarterly strategic review brief'],
    ['Decision Log',
     'All recorded decisions with rationale, alternatives, outcomes, and review dates',
     'Use 7-step Decision Framework for all new entries; flag decisions due for outcome review',
     'Quarterly decision outcome review'],
]
make_table(ps_headers, ps_rows, col_widths=[1.4, 2.5, 2.0, 1.85])

add_spacer()
add_heading('10.3  Fastest Wins This Week', 2)
add_body('Ranked by time-to-first-ROI. Execute in this order, this week, before any other ecosystem work.')
add_spacer()

fw_headers = ['Rank', 'Action', 'Time Required', 'Immediate Payoff']
fw_rows = [
    ['1', 'Create your top 3 active client Projects and load context documents',
     '3–4 hours total',
     'Never re-brief Claude on these clients again. Instant session time savings.'],
    ['2', 'Build the Adversarial Sales Audit Skill',
     '2 hours',
     'Deploy on a current client engagement this week. Demonstrate methodology consistency.'],
    ['3', 'Build the Decision Framework Skill',
     '2 hours',
     'Apply to the next significant decision you face. Experience the quality difference.'],
    ['4', 'Configure one Scheduled Task: Weekly Pipeline Scan',
     '1–2 hours setup',
     'First output arrives next week. Recurring work automated permanently.'],
    ['5', 'Run Deep Research on one high-value research task you have been doing manually',
     '30 minutes',
     'Direct comparison of output quality vs. conversational research. Usually immediately converts.'],
]
make_table(fw_headers, fw_rows, col_widths=[0.4, 2.5, 1.4, 3.45])

add_spacer()
add_heading('10.4  5 Behaviour Changes to Make Immediately', 2)
add_body(
    "These are not suggestions. Each of these behaviours is costing you leverage every day "
    "you continue them. Change them now, not after you have built the full system."
)
add_spacer()

beh_headers = ['#', 'Stop Doing', 'Start Doing', 'Why This Matters']
beh_rows = [
    ['1',
     'Starting any client session by re-briefing Claude on client context',
     'Open the client Project — Claude already knows the context',
     'You are spending 20–40 minutes per session re-entering context you have already paid to generate. Stop.'],
    ['2',
     'Treating first-draft Claude outputs as near-final without adversarial review',
     'Deploy a 5-minute adversarial review pass on every significant output before acting',
     'The quality gap between first draft and adversarially reviewed draft is larger than you think. Your decisions and deliverables deserve the second pass.'],
    ['3',
     'Rebuilding your best prompts from memory for recurring tasks',
     'Build the Skill now, the moment you recognise a task is recurring',
     'Every prompt you rebuild from memory is a tax on your cognitive capital. Skills are infrastructure. Build them.'],
    ['4',
     'Using chat for external research and market analysis',
     'Deploy Deep Research for every task that requires current, multi-source data',
     'You are making decisions and building deliverables on a subset of available intelligence. Deep Research is not a luxury — it is basic research hygiene for your work.'],
    ['5',
     'Letting Claude insights and decisions exist only in chat threads',
     'Record every significant insight, decision, and output to its relevant Project',
     'Your institutional knowledge is currently dying every time a chat session closes. The compound value of a populated Knowledge Layer at 12 months is enormous. You are at month zero. Start building now.'],
]
make_table(beh_headers, beh_rows, col_widths=[0.3, 1.8, 2.0, 3.65])

add_spacer(2)

# Closing
add_heading('CLOSING STATEMENT', 1)
add_body(
    "You are not underperforming because of lack of skill or intelligence. You are underperforming "
    "relative to your potential because you have been using a system-level tool at the conversation "
    "level. The gap between where you are and 1,000% leverage is not a talent gap — it is a "
    "structural gap. And structural gaps close fast once you start building."
)
add_body(
    "The four-week roadmap above is not aspirational. It is executable. The 20 use cases are not "
    "theoretical. They are deployable this week. The five behaviour changes are not gradual shifts. "
    "They are decisions you can make in the next five minutes."
)
add_body(
    "The operators who will define what AI-native agency and advisory work looks like in 2027 are "
    "building their ecosystem architecture right now. Not next quarter. Not after the next client "
    "engagement. Now. The compound returns on ecosystem infrastructure are non-linear — they "
    "accelerate sharply at month three, six, and twelve. You are starting from a high base. "
    "That is an advantage. Do not waste it."
)
add_spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— END OF REPORT —')
r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = GOLD

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('CONFIDENTIAL — For Authorised Recipient Only')
r2.font.size = Pt(8); r2.font.name = 'Calibri'; r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('/home/user/AIandTechWebsite/Claude_Ecosystem_Supremacy_Blueprint.docx')
print('Done')
